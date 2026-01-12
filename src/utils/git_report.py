from docx import Document

# Create a Word document
doc = Document()
doc.add_heading("Exercise Report: Try out Git", level=1)

doc.add_heading("Objective", level=2)
doc.add_paragraph("To install, configure, and initialize Git in Azure Cloud Shell, and explore basic Git commands.")

doc.add_heading("Step 1: Verify Git Installation", level=2)
doc.add_paragraph("Check if Git is installed in Cloud Shell:")
doc.add_paragraph("git --version", style="Quote")
doc.add_paragraph("Expected Output:\ngit version 2.x.x\n(Example: git version 2.7.4)")

doc.add_heading("Step 2: Configure Git Username and Email", level=2)
doc.add_paragraph("Set the username and email to identify commits.")
doc.add_paragraph('git config --global user.name "<USER_NAME>"\n'
                  'git config --global user.email "<USER_EMAIL>"\n'
                  'git config --list', style="Quote")
doc.add_paragraph("Expected Output:\nuser.name=<USER_NAME>\nuser.email=<USER_EMAIL>")

doc.add_heading("Step 3: Create a Project Directory", level=2)
doc.add_paragraph("Create and move into a project folder:")
doc.add_paragraph("mkdir Cats\ncd Cats", style="Quote")

doc.add_heading("Step 4: Initialize a Git Repository", level=2)
doc.add_paragraph("Initialize the repository and set the default branch to main:")
doc.add_paragraph("git init --initial-branch=main\nor\ngit init -b main", style="Quote")
doc.add_paragraph("For earlier versions of Git:")
doc.add_paragraph("git init\ngit checkout -b main", style="Quote")
doc.add_paragraph("Expected Output:\nInitialized empty Git repository in /home/<user>/Cats/.git/\n"
                  "Switched to a new branch 'main'")

doc.add_heading("Step 5: Check Repository Status", level=2)
doc.add_paragraph("Check the current branch and status:")
doc.add_paragraph("git status", style="Quote")
doc.add_paragraph("Expected Output:\nOn branch main\nNo commits yet\nnothing to commit (create/copy files and use \"git add\" to track)")

doc.add_heading("Step 6: View Repository Contents", level=2)
doc.add_paragraph("List all files including hidden ones:")
doc.add_paragraph("ls -a", style="Quote")
doc.add_paragraph("Expected Output:\n.  ..  .git\n\nThe .git directory stores Git metadata and history. Do not modify it manually.")

doc.add_heading("Step 7: Explore Git Help", level=2)
doc.add_paragraph("Explore available Git commands:")
doc.add_paragraph("git --help", style="Quote")
doc.add_paragraph("To get detailed help for a specific command:")
doc.add_paragraph("git help <command>", style="Quote")

doc.add_heading("Summary", level=2)
doc.add_paragraph("• Verified Git installation in Azure Cloud Shell\n"
                  "• Configured Git username and email\n"
                  "• Created a project directory named 'Cats'\n"
                  "• Initialized a Git repository with a main branch\n"
                  "• Checked Git status and confirmed .git directory creation\n"
                  "• Used Git help to explore available commands")

doc.add_heading("Outcome", level=2)
doc.add_paragraph("Successfully set up Git, configured user identity, and created the first local repository, preparing for future Git operations such as adding files, committing, and pushing to remote repositories.")

# Save the Word document
doc.save("Try_out_Git_Report.docx")
print("✅ Word report created successfully: Try_out_Git_Report.docx")
